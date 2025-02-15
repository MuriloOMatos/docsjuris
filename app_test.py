from flask import Flask, request, send_file, render_template
from docx import Document
import io

app = Flask(__name__)

# Função para calcular a diferença entre os juros do contrato e a taxa média do BACEN
def calcular_diferenca(valor_financiado, taxa_contrato, taxa_media, parcelas):
    juros_contrato = (float(taxa_contrato) / 100) * float(valor_financiado) * int(parcelas)
    juros_media = (float(taxa_media) / 100) * float(valor_financiado) * int(parcelas)
    return juros_contrato - juros_media

@app.route('/')
def index():
    return render_template('index.html')  # Certifique-se de que o arquivo index.html esteja na pasta "templates"

@app.route('/gerar-peticao', methods=['POST'])
def gerar_peticao():
    # Coleta dos dados gerais do formulário
    modelo_selecionado = request.form['modelo_peticao']  # Ex: "modelo_1.docx", "modelo_2.docx" ou "modelo_3.docx"
    dados = {
        'renda_mensal': request.form['renda_mensal'],
        'data_contratacao': request.form['data_contratacao'],
        'parcela_pessoal': request.form['parcela_pessoal'],
        'valor_financiado': request.form['valor_financiado'],
        'quantidade_parcelas': request.form['quantidade_parcelas'],
        'taxa_juros_contrato': request.form['taxa_juros_contrato'],
        'taxa_media_bacen': request.form['taxa_media_bacen']
    }

    # Conversão para numérico para os cálculos
    try:
        renda_mensal = float(request.form['renda_mensal'])
    except:
        renda_mensal = 0.0
    try:
        parcela_pessoal = float(request.form['parcela_pessoal'])
    except:
        parcela_pessoal = 0.0
    try:
        valor_fin = float(request.form['valor_financiado'])
    except:
        valor_fin = 0.0
    try:
        num_parcelas = int(request.form['quantidade_parcelas'])
    except:
        num_parcelas = 0

    # Se houver um campo global "parcela_consignado" (caso esteja definido fora dos grupos de empréstimos)
    try:
        parcela_consignado_global = float(request.form.get('parcela_consignado', 0))
    except:
        parcela_consignado_global = 0.0

    # Cálculo do valor líquido disponível (Renda - (Parcela Consignada + Parcela Pessoal))
    valor_liquido = renda_mensal - (parcela_consignado_global + parcela_pessoal)
    dados['valor_liquido'] = f"R$ {valor_liquido:.2f}"

    # Cálculo do comprometimento da renda (global)
    comprometimento = ((parcela_consignado_global + parcela_pessoal) / renda_mensal * 100) if renda_mensal != 0 else 0
    dados['comprometimento_renda'] = f"{comprometimento:.2f}%"

    # Cálculo do valor total do empréstimo (global)
    total_emprestimo = parcela_consignado_global * num_parcelas
    dados['total_emprestimo'] = f"R$ {total_emprestimo:.2f}"

    # Cálculo do montante acumulado com juros compostos (global)
    taxa_juros = float(request.form['taxa_juros_contrato']) / 100  # taxa em decimal
    acima_do_financiado = valor_fin * ((1 + taxa_juros) ** num_parcelas)
    dados['acima_do_financiado'] = f"R$ {acima_do_financiado:.2f}"

    # Cálculo da diferença entre os juros do contrato e a taxa média do BACEN (global)
    diferenca = calcular_diferenca(request.form['valor_financiado'],
                                   request.form['taxa_juros_contrato'],
                                   request.form['taxa_media_bacen'],
                                   request.form['quantidade_parcelas'])
    dados['diferenca_contrato1'] = f"R$ {diferenca:.2f}"

    # Coleta dos dados dinâmicos dos empréstimos 2 e 3 (ou quantos forem informados)
    try:
        num_loans = int(request.form['modelo_peticao'])
    except (TypeError, ValueError):
        num_loans = 0

    # Itera de 0 até num_loans - 1 para coletar dados de cada empréstimo
    for i in range(num_loans):
        valor = request.form.get(f"emprestimos[{i}][valor]", "")
        parcela_consignada = request.form.get(f"emprestimos[{i}][parcela_consignada]", "")
        data = request.form.get(f"emprestimos[{i}][data]", "")
        parcelas = request.form.get(f"emprestimos[{i}][parcelas]", "")
        taxa = request.form.get(f"emprestimos[{i}][taxa]", "")
        
        dados[f"emprestimos_{i}_valor"] = valor
        dados[f"emprestimos_{i}_parcela_consignada"] = parcela_consignada
        dados[f"emprestimos_{i}_data"] = data
        dados[f"emprestimos_{i}_parcelas"] = parcelas
        dados[f"emprestimos_{i}_taxa"] = taxa

    # Abre o modelo Word selecionado
    doc = Document(modelo_selecionado)

    # Substitui os placeholders nos parágrafos do documento
    for paragraph in doc.paragraphs:
        for key, value in dados.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
    
    # Substitui os placeholders em tabelas, se houver
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in dados.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
    
    # Salva o documento gerado em memória
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(output, download_name="acao_revisional.docx", as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
