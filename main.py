from flask import Flask, render_template, request, send_file, abort, redirect, url_for, session
from docx import Document
import io
import os
import requests
from datetime import datetime
from decimal import Decimal, InvalidOperation
import bleach
from calendar import monthrange
from functools import lru_cache
from requests.adapters import HTTPAdapter
from requests.packages.urllib3.util.retry import Retry
from werkzeug.security import generate_password_hash, check_password_hash
import logging
import zipfile
import tempfile
import psycopg2
from psycopg2.extras import RealDictCursor
import json
import re

# Lista de templates válidos - atualizada com os novos templates
VALID_TEMPLATES = [
    'declaracao_hiposuficiencia',
    'contratos_honorarios',
    'declaracao_contrato_digital',
    'declaracao_procuradores',
    'declaracao_ir',
    'procuracao',
    'declaracao_residencia'
]

# Adicionar templates dinamicamente para autor e réu
for foro in ['autor', 'reu']:
    for i in range(1, 4):  # 1 a 3 empréstimos
        VALID_TEMPLATES.append(f'Contrato_Foro_{foro.capitalize()}_{i}')

# Mapeamento de valores para exibição
MAPEAMENTO_VALORES = {
    # Tipos de petição
    "judicial": "Judicial",
    "extrajudicial": "Extrajudicial",
    "administrativa": "Administrativa",
    
    # Foro
    "autor": "Autor",
    "reu": "Réu",
    
    # Empréstimos consignados
    "sim": "Sim",
    "nao": "Não",
    
    # Fontes de renda
    "salario": "Salário",
    "beneficio_deficiencia": "Benefício de Prestação Continuada à Pessoa com Deficiência",
    "beneficio_idosa": "Benefício de Prestação Continuada à Pessoa Idosa",
    "pensao_morte": "Benefício de Pensão Por Morte Previdenciária",
    "aposentadoria_idade": "Aposentadoria Por Idade Previdenciária",
    "aposentadoria_tempo": "Aposentadoria Por Tempo de Contribuição Previdenciária",
    "aposentadoria_incapacidade": "Aposentadoria Por Incapacidade Permanente Previdenciária",
    "maes_pernambuco": "Mães de Pernambuco",
    "auxilio_gas": "Auxílio Gás dos Brasileiros",
    "bolsa_familia": "Bolsa Família",
    "auxilio_incapacidade": "Benefício de Auxílio Por Incapacidade Temporária Previdenciário",
    "auxilio_acidente": "Benefício de Auxílio-Acidente",
    
    # Conjunto probatório
    "declaracao_hipossuficiencia": "Declaração de Hipossuficiência firmada pela parte autora",
    "isencao_imposto_renda": "Declaração de isenção de Imposto de Renda",
    "print_receita_federal": "Print da Receita Federal demonstrando inexistência de dados quanto à Declaração Anual de Imposto de Renda",
    "extratos_inss": "Extratos de benefício previdenciário (INSS)",
    "extratos_bancarios": "Extratos bancários atualizados",
    "ctps_digital": "Carteira de Trabalho e Previdência Social (CTPS) Digital",
    "cadastro_unico": "Folha resumo do Cadastro Único para Programas Sociais do Governo Federal",
    
    # Estados
    "AC": "Acre",
    "AL": "Alagoas",
    "AP": "Amapá",
    "AM": "Amazonas",
    "BA": "Bahia",
    "CE": "Ceará",
    "DF": "Distrito Federal",
    "ES": "Espírito Santo",
    "GO": "Goiás",
    "MA": "Maranhão",
    "MT": "Mato Grosso",
    "MS": "Mato Grosso do Sul",
    "MG": "Minas Gerais",
    "PA": "Pará",
    "PB": "Paraíba",
    "PR": "Paraná",
    "PE": "Pernambuco",
    "PI": "Piauí",
    "RJ": "Rio de Janeiro",
    "RN": "Rio Grande do Norte",
    "RS": "Rio Grande do Sul",
    "RO": "Rondônia",
    "RR": "Roraima",
    "SC": "Santa Catarina",
    "SP": "São Paulo",
    "SE": "Sergipe",
    "TO": "Tocantins",
    
    # Novos campos
    "estado_comarca": "Estado da Comarca",
    "cidade_comarca": "Cidade da Comarca", 
    "estado_oab": "Estado da OAB",
    "numero_oab": "Número da OAB",
}

# Configuração de logging
logging.basicConfig(level=logging.DEBUG)

# Inicialização do Flask
app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', os.urandom(24).hex())

# Usuários fictícios
USERS = {
    'GMadvogados': generate_password_hash('GM1252')
}

# Configuração via variável de ambiente
SERIES_BACEN = {'pessoal_fisica': int(os.getenv('SERIE_BACEN', 25464))}

# Sessão HTTP com retry
http_session = requests.Session()
retries = Retry(total=3, backoff_factor=0.3, status_forcelist=[500, 502, 503, 504], allowed_methods=["GET"])
adapter = HTTPAdapter(max_retries=retries)
http_session.mount("http://", adapter)
http_session.mount("https://", adapter)

def get_db_connection():
    try:
        return psycopg2.connect(
            host=os.getenv("DB_HOST"),
            database=os.getenv("DB_NAME"),
            user=os.getenv("DB_USER"),
            password=os.getenv("DB_PASSWORD"),
            port=os.getenv("DB_PORT"),
            cursor_factory=RealDictCursor
        )
    except psycopg2.Error as e:
        app.logger.error(f"Erro ao conectar ao banco de dados: {str(e)}")
        raise

@app.route('/login', methods=['GET', 'POST'])
def login():
    app.logger.debug("Acessando rota /login")
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        app.logger.debug(f"Tentativa de login com usuário: {username}")
        
        if username in USERS and check_password_hash(USERS[username], password):
            session['logged_in'] = True
            session['username'] = username
            app.logger.debug("Login bem-sucedido")
            return redirect(url_for('index'))
        else:
            app.logger.debug("Falha no login")
            return render_template('login.html', error='Usuário or senha inválidos')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    app.logger.debug("Executando logout")
    session.pop('logged_in', None)
    session.pop('username', None)
    return redirect(url_for('login'))

def login_required(f):
    def wrap(*args, **kwargs):
        if 'logged_in' not in session:
            app.logger.debug("Usuário não autenticado, redirecionando para login")
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    wrap.__name__ = f.__name__
    return wrap

def _obter_dados_api(url, params):
    headers = {'User-Agent': 'Python/AppRevisaoContratos'}
    try:
        response = http_session.get(url, params=params, timeout=10, headers=headers)
        response.raise_for_status()
        dados = response.json()
        if not dados or 'valor' not in dados[0]:
            raise ValueError("Nenhum dado válido retornado pela API")
        return dados
    except Exception as e:
        app.logger.error(f"Erro ao acessar API BACEN: {str(e)}")
        raise

@lru_cache(maxsize=128)
def get_bacen_taxa_historico(data_emprestimo):
    try:
        codigo_serie = SERIES_BACEN['pessoal_fisica']
        mes_ano = data_emprestimo.strftime("%m/%Y")
        _, last_day = monthrange(data_emprestimo.year, data_emprestimo.month)
        url = f"https://api.bcb.gov.br/dados/serie/bcdata.sgs.{codigo_serie}/dados"
        params = {
            'formato': 'json',
            'dataInicial': f"01/{mes_ano}",
            'dataFinal': f"{last_day:02d}/{mes_ano}"
        }
        dados = _obter_dados_api(url, params)
        return float(dados[0]['valor'])
    except Exception as e:
        app.logger.error(f"Erro ao buscar taxa histórica BACEN: {str(e)}")
        return 0.01

def get_bacen_taxa_atual():
    try:
        codigo_serie = SERIES_BACEN['pessoal_fisica']
        url = f"https://api.bcb.gov.br/dados/serie/bcdata.sgs.{codigo_serie}/dados/ultimos/1"
        params = {'formato': 'json'}
        dados = _obter_dados_api(url, params)
        return float(dados[0]['valor'])
    except Exception as e:
        app.logger.error(f"Erro ao buscar taxa atual BACEN: {str(e)}")
        return 0.01

@app.route("/peticoes")
@login_required
def peticoes():
    return render_template("peticoes.html")

@app.route('/bancos')
@login_required
def listar_bancos():
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT * FROM bancos ORDER BY nome_banco ASC")
        bancos = cur.fetchall()
        cur.close()
        conn.close()
        return render_template('bancos.html', bancos=bancos)
    except psycopg2.Error as e:
        app.logger.error(f"Erro ao acessar bancos: {str(e)}")
        return render_template('bancos.html', bancos=[], error="Erro ao carregar bancos")

@app.route('/bancos/adicionar', methods=['POST'])
@login_required
def adicionar_banco():
    codigo = request.form.get('codigo_banco')
    nome = request.form.get('nome_banco')

    if not codigo or not nome:
        return "Código e Nome do banco são obrigatórios.", 400

    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT 1 FROM bancos WHERE codigo_banco = %s", (codigo,))
        if cur.fetchone():
            cur.close()
            conn.close()
            return "Código do banco já existe.", 400

        cur.execute(
            "INSERT INTO bancos (codigo_banco, nome_banco) VALUES (%s, %s)",
            (codigo, nome)
        )
        conn.commit()
        cur.close()
        conn.close()
        return redirect(url_for('listar_bancos'))
    except psycopg2.Error as e:
        app.logger.error(f"Erro ao adicionar banco: {str(e)}")
        return "Erro ao adicionar banco.", 500

@app.route('/documentos')
@login_required
def documentos():
    app.logger.debug("Acessando rota /documentos")
    
    bancos = [
        {"codigo_banco": "01", "nome_banco": "BANCO DO BRASIL.", "cnpj": "00.000.000/0001-00"},
        {"codigo_banco": "02", "nome_banco": "CAIXA ECONÔMICA.", "cnpj": "00.000.000/0001-91"},
        {"codigo_banco": "03", "nome_banco": "BANCO MERCANTIL DO BRASIL S.A.", "cnpj": "17.184.037/0001-10"},
        {"codigo_banco": "04", "nome_banco": "BANDO CREFISA S.A.", "cnpj": "61.033.106/0001-86"},
        {"codigo_banco": "05", "nome_banco": "BANCO BMG S.A.", "cnpj": "61.186.680/0001-74"},
        {"codigo_banco": "06", "nome_banco": "BANCO AGIBANK S.A.", "cnpj": "10.664.513/0001-50"},
        {"codigo_banco": "07", "nome_banco": "CREFAZ SOCIEDADE DE CREDITO AO MICROEMPREENDEDOR E A EMPRESA DE PEQUENO PORTE S.A.", "cnpj": "18.188.384/0001-83"},
        {"codigo_banco": "08", "nome_banco": "BANCO BRADESCO S.A.", "cnpj": "60.746.948/0001-12"}
    ]
    
    return render_template('documentos.html', bancos=bancos)

def calcular_diferenca(valor, taxa_contrato, taxa_media, parcelas):
    try:
        valor_dec = Decimal(str(valor))
        taxa_contrato_dec = Decimal(str(taxa_contrato))
        taxa_media_dec = Decimal(str(taxa_media))
        parcelas_dec = Decimal(str(parcelas))
        
        if any(x <= 0 for x in [valor_dec, taxa_contrato_dec, taxa_media_dec, parcelas_dec]):
            raise ValueError("Todos os valores devem ser positivos")
        
        diferenca = valor_dec * (taxa_contrato_dec - taxa_media_dec) * parcelas_dec / 100
        return float(diferenca)
    except (InvalidOperation, ValueError) as e:
        app.logger.error(f"Erro no cálculo de diferença: {str(e)}")
        raise ValueError(f"Erro ao calcular diferença: {str(e)}")

def validar_dados_entrada(form):
    required_fields = ['renda_mensal', 'parcela_pessoal', 'modelo_peticao', 'foro']
    for field in required_fields:
        if field not in form:
            raise ValueError(f"Campo obrigatório '{field}' ausente")
    
    try:
        renda_mensal = Decimal(form['renda_mensal'].replace(",", "."))
        parcela_pessoal = Decimal(form['parcela_pessoal'].replace(",", "."))
        if renda_mensal <= 0 or parcela_pessoal < 0:
            raise ValueError("Renda mensal deve ser positiva and parcela pessoal não pode ser negativa")
        
        num_emprestimos = int(form['modelo_peticao'])
        if num_emprestimos not in [1, 2, 3]:
            raise ValueError("Número de empréstimos inválido (deve be 1, 2 ou 3)")
        
        foro = form['foro'].lower()
        if foro not in ['autor', 'reu']:
            raise ValueError("Foro inválido (deve ser 'autor' ou 'reu')")
            
    except (ValueError, InvalidOperation) as e:
        raise ValueError(f"Erro na validação de dados: {str(e)}")
    
    return num_emprestimos, foro

def calculos_emprestimo(form, num_emprestimos):
    emprestimos = []
    total_consignado = Decimal('0')
    total_emprestimo_geral = Decimal('0')
    total_dobro_geral = Decimal('0')
    renda_mensal = Decimal(form['renda_mensal'].replace(",", "."))
    parcela_pessoal = Decimal(form['parcela_pessoal'].replace(",", "."))
    
    for i in range(num_emprestimos):
        prefix = f'emprestimos[{i}]'
        emp_data = form.get(f'{prefix}[data]')
        try:
            data_emprestimo = datetime.strptime(emp_data, '%d/%m/%Y')
            if data_emprestimo > datetime.now():
                raise ValueError(f"Data do empréstimo {i+1} não pode ser no futuro")
        except ValueError as e:
            raise ValueError(f"Data do empréstimo {i+1} inválida ou no futuro. Use DD/MM/YYYY") from e
        
        taxa_media = get_bacen_taxa_historico(data_emprestimo)
        if taxa_media is None:
            raise ValueError(f"Não foi possível obter a taxa média para o empréstimo {i+1} ({emp_data})")
        
        valor_str = form.get(f'{prefix}[valor]', '0').replace(",", ".")
        parcela_str = form.get(f'{prefix}[parcela_consignada]', '0').replace(",", ".")
        parcelas_str = form.get(f'{prefix}[parcelas]', '0')
        taxa_contrato_str = form.get(f'{prefix}[taxa]', '0').replace(",", ".")
        contrato_str = form.get(f'{prefix}[contrato]', 'N/A')  # Capturar número do contrato
        
        try:
            valor = Decimal(valor_str)
            parcela = Decimal(parcela_str)
            parcelas = Decimal(parcelas_str)
            taxa_contrato = Decimal(taxa_contrato_str)
            
            if not all(x > 0 for x in [valor, parcela, parcelas, taxa_contrato]):
                raise ValueError(f"Valores do empréstimo {i+1} devem ser positivos")
            
            if Decimal(str(taxa_media)) <= 0:
                raise ValueError(f"Taxa média BACEN deve ser positiva para empréstimo {i+1}")
        except (ValueError, InvalidOperation) as e:
            raise ValueError(f"Erro nos valores do empréstimo {i+1}: {str(e)}")
        
        total_emprestimo = parcela * parcelas
        total_emprestimo_geral += total_emprestimo
        def_emprestimos = total_emprestimo / valor if valor != 0 else Decimal('0')
        taxa_media_dec = Decimal(str(taxa_media)) / 100
        parcela_pessoal_atual = valor * (taxa_media_dec / (1 - (1 + taxa_media_dec) ** -parcelas)) if (1 + taxa_media_dec) ** -parcelas != 1 else Decimal('0')
        total_emprestimo_bacen = parcela_pessoal_atual * parcelas
        dif_bacen = valor / parcela_pessoal_atual if parcela_pessoal_atual != 0 else Decimal('0')
        taxa_contrato_dec = taxa_contrato / 100
        vlr_total_emprestimo1 = valor * (1 + Decimal(str(taxa_media)) / 100) ** parcelas
        vlr_total_emprestimo2 = valor * (1 + taxa_contrato_dec) ** parcelas
        org_bacen = abs(total_emprestimo - total_emprestimo_bacen)
        org_div = (total_emprestimo_geral / vlr_total_emprestimo1) if vlr_total_emprestimo1 != 0 else Decimal('0')
        total_dobro = org_bacen * 2
        total_dobro_geral += total_dobro
        comprometimento_renda = parcela + parcela_pessoal
        comprometimento_porcentagem = (comprometimento_renda / renda_mensal * 100) if renda_mensal != 0 else Decimal('0')
        renda_atual = renda_mensal - parcela_pessoal - parcela
        
        emprestimo = {
            'data': emp_data,
            'valor': valor_str,
            'parcela': parcela_str,
            'parcelas': parcelas_str,
            'taxa': taxa_contrato_str,
            'contrato': contrato_str,  # Adicionar número do contrato
            'taxa_media': f"{taxa_media:.2f}",
            'diferenca': f"{calcular_diferenca(valor, taxa_contrato, taxa_media, parcelas):.2f}",
            'total_emprestimo': f"{total_emprestimo:.2f}",
            'def_emprestimos': f"{def_emprestimos:.2f}",
            'parcela_pessoal_atual': f"{parcela_pessoal_atual:.2f}",
            'dif_bacen': f"{dif_bacen:.2f}",
            'vlr_total_emprestimo1': f"{vlr_total_emprestimo1:.2f}",
            'vlr_total_emprestimo2': f"{vlr_total_emprestimo2:.2f}",
            'org_bacen': f"{org_bacen:.2f}",
            'org_div': f"{org_div:.2f}",
            'total_dobro': f"{total_dobro:.2f}",
            'valor_causa': f"{Decimal('5000') + total_dobro:.2f}",
            'dadovalorcausa': f"0.00",
            'comprometimento_renda': f"{comprometimento_renda:.2f}",
            'comprometimento_porcentagem': f"{comprometimento_porcentagem:.2f}",
            'renda_atual': f"{renda_atual:.2f}",
            'total_emprestimo_bacen': f"{total_emprestimo_bacen:.2f}",
        }
        
        emprestimos.append(emprestimo)
        total_consignado += parcela
        
    dadovalorcausa = total_dobro_geral + Decimal('5000')
    valor_causa = dadovalorcausa
    
    for emp in emprestimos:
        emp['dadovalorcausa'] = f"{dadovalorcausa:.2f}"
    
    return emprestimos, total_consignado, total_emprestimo_geral, def_emprestimos, parcela_pessoal_atual, dif_bacen, vlr_total_emprestimo1, vlr_total_emprestimo2, org_bacen, org_div, total_dobro, valor_causa, comprometimento_renda, renda_atual, comprometimento_porcentagem, total_emprestimo_bacen, total_dobro_geral, dadovalorcausa

def determinar_template_peticao(foro, num_emprestimos):
    """Determina qual template usar baseado no foro e número de empréstimos"""
    template_name = f"Contrato_Foro_{foro.capitalize()}_{num_emprestimos}"
    
    # Verificar se o template existe
    template_path = os.path.join('modelos', f"{template_name}.docx")
    if not os.path.exists(template_path):
        # Fallback para template genérico se o específico não existir
        app.logger.warning(f"Template {template_name} não encontrado, usando fallback")
        template_name = f"Contrato_Foro_{foro.capitalize()}_1"
        template_path = os.path.join('modelos', f"{template_name}.docx")
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Nenhum template válido encontrado para foro {foro}")
    
    return template_name

def gerar_documento(dados, num_emprestimos, foro='autor'):
    """Gera documento usando o template correto baseado no foro e número de empréstimos"""
    template_name = determinar_template_peticao(foro, num_emprestimos)
    template_path = os.path.abspath(os.path.join("modelos", f"{template_name}.docx"))
    
    if not os.path.exists(template_path):
        app.logger.error(f"Template {template_path} não encontrado")
        raise FileNotFoundError(f"Modelo de documento {template_name}.docx não encontrado")
    
    doc = Document(template_path)
    
    # DEBUG: Log dos dados recebidos
    app.logger.debug(f"Dados recebidos para placeholders: {list(dados.keys())}")
    
    # Mapeamento completo de placeholders - INCLUINDO BENEFICIOS
    replacements = {
        # Dados básicos da petição
        'foro': dados.get('foro', 'Autor'),
        'tipo_peticao': dados.get('tipo_peticao', ''),
        'banco': dados.get('banco', 'Banco Não Especificado'),
        'possui_emprestimos': dados.get('possui_emprestimos', ''),
        'fontes_renda': dados.get('fontes_renda', 'Nenhuma fonte de renda selecionada'),
        'conjunto_probatorio': dados.get('conjunto_probatorio', 'Nenhum documento probatório selecionado'),
        'estado_comarca': dados.get('estado_comarca', ''),
        'cidade_comarca': dados.get('cidade_comarca', ''),
        'advogado': dados.get('advogado', ''),
        'estado_oab': dados.get('estado_oab', ''),
        'numero_oab': dados.get('numero_oab', ''),
        'beneficios': dados.get('beneficios', dados.get('fontes_renda', 'Nenhuma fonte de renda selecionada')),  # NOVO CAMPO
        'BENEFICIOS': dados.get('beneficios', dados.get('fontes_renda', 'Nenhuma fonte de renda selecionada')),  # NOVO CAMPO em maiúsculas
        
        # Dados financeiros
        'renda_mensal': dados.get('renda_mensal', ''),
        'parcela_pessoal': dados.get('parcela_pessoal', ''),
        'valor_liquido': dados.get('valor_liquido', ''),
        'comprometimento': dados.get('comprometimento', ''),
        'diario': dados.get('diario', ''),
        'comprometimento_porcentagem': dados.get('comprometimento_porcentagem', ''),
        
        # Valores calculados
        'valor_causa': dados.get('valor_causa', ''),
        'total_dobro_geral': dados.get('total_dobro_geral', ''),
        'renda_atual': dados.get('renda_atual', ''),
        'total_emprestimo': dados.get('total_emprestimo', ''),
        'def_emprestimos': dados.get('def_emprestimos', ''),
        'parcela_pessoal_atual': dados.get('parcela_pessoal_atual', ''),
        'dif_bacen': dados.get('dif_bacen', ''),
        'vlr_total_emprestimo1': dados.get('vlr_total_emprestimo1', ''),
        'vlr_total_emprestimo2': dados.get('vlr_total_emprestimo2', ''),
        'org_bacen': dados.get('org_bacen', ''),
        'org_div': dados.get('org_div', ''),
        'total_dobro': dados.get('total_dobro', ''),
        'dadovalorcausa': dados.get('dadovalorcausa', ''),
        'total_emprestimo_bacen': dados.get('total_emprestimo_bacen', ''),
        
        # Placeholders específicos do template - ADICIONAR VARIANTES
        'comarca': f"{dados.get('cidade_comarca', '')}/{dados.get('estado_comarca', '')}",
        'cidade': dados.get('cidade_comarca', ''),
        'estado': dados.get('estado_comarca', ''),
        'beneficio_recebido': 'Benefício Previdenciário',
        'numero_contrato': dados.get('numero_contrato', 'N/A'),
        'data': datetime.now().strftime('%d/%m/%Y'),
        'n_oab': dados.get('numero_oab', '1252'),
        'cnpj_banco': '00.000.000/0001-00',
        'endereco_banco': 'Endereço não especificado',
    }
    
    # Adicionar dados dos empréstimos se existirem
    if 'emprestimos' in dados and isinstance(dados['emprestimos'], list):
        for i, emp in enumerate(dados['emprestimos']):
            for key, value in emp.items():
                # Formatar para o padrão do template: emprestimos_0_*
                replacements[f'emprestimos_{i}_{key}'] = value
                
                # Adicionar também placeholders específicos para número de contrato
                if key == 'contrato':
                    replacements[f'numero_contrato_{i}'] = value
    
    # DEBUG: Log dos replacements
    app.logger.debug(f"Replacements preparados: {list(replacements.keys())}")
    
    # Dicionário de trechos que devem estar em negrito
    bold_sections = {
        "CONTRATADO: GUILHERME ESTEVES DOS SANTOS MORAES": True,
        "AÇÃO REVISIONAL DE CONTRATO DE EMPRÉSTIMO PESSOAL NÃO CONSIGNADO": True
    }
    
    # Função para substituir placeholders em texto - MELHORADA
    def substituir_placeholders(texto):
        if not texto:
            return texto
            
        novo_texto = texto
        for key, value in replacements.items():
            # Tentar múltiplos formatos de placeholders
            placeholders_to_try = [
                f'{{{{{key}}}}}',           # {{key}}
                f'{{{{{key.upper()}}}}}',   # {{KEY}}
                f'{{{{{key.lower()}}}}}',   # {{key}} (minúsculo)
            ]
            
            for placeholder in placeholders_to_try:
                if placeholder in novo_texto:
                    novo_texto = novo_texto.replace(placeholder, str(value))
                    app.logger.debug(f"Substituído {placeholder} por {value}")
        
        return novo_texto
    
    # Substituir placeholders em parágrafos
    for p in doc.paragraphs:
        original_text = p.text
        new_text = substituir_placeholders(original_text)
        
        # Aplicar negrito nos trechos específicos
        for text, should_bold in bold_sections.items():
            if text in new_text and should_bold:
                for run in p.runs:
                    if text in run.text:
                        run.bold = True
        
        # Atualizar o texto do parágrafo apenas se houve mudanças
        if new_text != original_text:
            p.text = new_text
    
    # Substituir placeholders em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                new_text = substituir_placeholders(original_text)
                
                # Aplicar negrito nos trechos específicos
                for text, should_bold in bold_sections.items():
                    if text in new_text and should_bold:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if text in run.text:
                                    run.bold = True
                
                # Atualizar o texto da célula apenas se houve mudanças
                if new_text != original_text:
                    cell.text = new_text
    
    # Substituir placeholders em cabeçalhos e rodapés
    for section in doc.sections:
        # Cabeçalho
        if section.header:
            for p in section.header.paragraphs:
                original_text = p.text
                new_text = substituir_placeholders(original_text)
                if new_text != original_text:
                    p.text = new_text
        
        # Rodapé
        if section.footer:
            for p in section.footer.paragraphs:
                original_text = p.text
                new_text = substituir_placeholders(original_text)
                if new_text != original_text:
                    p.text = new_text
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route('/')
@login_required
def index():
    app.logger.debug("Acessando rota /")
    try:
        taxa_atual = get_bacen_taxa_atual()
        taxa_atual_display = f"{taxa_atual:.2f}%" if taxa_atual is not None else "Indisponível"
        bacen_data = {
            'taxa': taxa_atual_display,
            'data_atualizacao': datetime.now().strftime('%d/%m/%Y')
        }
        return render_template('index.html', **bacen_data)
    except Exception as e:
        app.logger.error(f"Erro na rota index: {str(e)}")
        return abort(500, "Erro interno ao carregar a página inicial")

def format_brl(valor):
    try:
        valor_dec = Decimal(str(valor))
        s = f"{valor_dec:,.2f}"
        s = s.replace(",", "X").replace(".", ",").replace("X", ".")
        return f"R$ {s}"
    except Exception:
        return str(valor)

def formatar_lista_selecionados(valores, tipo):
    """Formata uma lista de valores selecionados para exibição"""
    if not valores:
        return "Nenhum selecionado"
    
    if tipo == "array":
        # Se já é uma lista, apenas mapear os valores
        return ", ".join([MAPEAMENTO_VALORES.get(v, v) for v in valores])
    else:
        # Se é uma string JSON, converter para lista primeiro
        try:
            valores_lista = json.loads(valores)
            return ", ".join([MAPEAMENTO_VALORES.get(v, v) for v in valores_lista])
        except:
            return valores

@app.route('/gerar-peticao', methods=['POST'])
@login_required
def gerar_peticao():
    try:
        app.logger.debug("Iniciando geração de petição")
        
        # Validar e obter dados básicos
        num_emprestimos, foro = validar_dados_entrada(request.form)
        app.logger.debug(f"Foro selecionado: {foro}, Número de empréstimos: {num_emprestimos}")
        
        # Coletar dados da petição (primeira etapa) - INCLUIR BENEFICIOS
        dados_peticao = {
            'renda_mensal': request.form['renda_mensal'].replace(",", "."),
            'parcela_pessoal': request.form['parcela_pessoal'].replace(",", "."),
            'foro': MAPEAMENTO_VALORES.get(request.form.get('foro', 'autor'), 'Autor'),
            'tipo_peticao': MAPEAMENTO_VALORES.get(request.form.get('tipo_peticao', ''), ''),
            'banco': request.form.get('banco', ''),
            'possui_emprestimos': MAPEAMENTO_VALORES.get(request.form.get('possui_emprestimos', ''), ''),
            'fontes_renda': formatar_lista_selecionados(request.form.getlist('fontes_renda[]'), "array"),
            'conjunto_probatorio': formatar_lista_selecionados(request.form.getlist('conjunto_probatorio[]'), "array"),
            'estado_comarca': MAPEAMENTO_VALORES.get(request.form.get('estado_comarca', ''), request.form.get('estado_comarca', '')),
            'cidade_comarca': request.form.get('cidade_comarca', ''),
            'advogado': request.form.get('advogado', ''),
            'estado_oab': request.form.get('estado_oab', ''),
            'numero_oab': request.form.get('numero_oab', ''),
            'numero_contrato': request.form.get('numero_contrato', 'N/A'),
            # ADICIONAR O NOVO CAMPO BENEFICIOS
            'beneficios': request.form.get('beneficios', ''),
        }
        
        # DEBUG: Log dos dados da petição
        app.logger.debug(f"Dados da petição coletados: {dados_peticao}")
        
        # Calcular dados dos empréstimos
        emprestimos, total_consignado, total_emprestimo_geral, def_emprestimos, parcela_pessoal_atual, dif_bacen, vlr_total_emprestimo1, vlr_total_emprestimo2, org_bacen, org_div, total_dobro, valor_causa, comprometimento_renda, renda_atual, comprometimento_porcentagem, total_emprestimo_bacen, total_dobro_geral, dadovalorcausa = calculos_emprestimo(request.form, num_emprestimos)
        
        # Formatar valores para exibição
        for emp in emprestimos:
            emp['valor'] = format_brl(emp['valor'])
            emp['parcela'] = format_brl(emp['parcela'])
            emp['diferenca'] = format_brl(emp['diferenca'])
            emp['total_emprestimo'] = format_brl(emp['total_emprestimo'])
            emp['parcela_pessoal_atual'] = format_brl(emp['parcela_pessoal_atual'])
            emp['dif_bacen'] = format_brl(emp['dif_bacen'])
            emp['vlr_total_emprestimo1'] = format_brl(emp['vlr_total_emprestimo1'])
            emp['vlr_total_emprestimo2'] = format_brl(emp['vlr_total_emprestimo2'])
            emp['org_bacen'] = format_brl(emp['org_bacen'])
            emp['total_dobro'] = format_brl(emp['total_dobro'])
            emp['valor_causa'] = format_brl(emp['valor_causa'])
            emp['renda_atual'] = format_brl(emp['renda_atual'])
            emp['total_emprestimo_bacen'] = format_brl(emp['total_emprestimo_bacen'])
            emp['dadovalorcausa'] = format_brl(emp['dadovalorcausa'])
            
        dados_peticao['emprestimos'] = emprestimos
        renda = Decimal(dados_peticao['renda_mensal'])
        parcela_pessoal = Decimal(dados_peticao['parcela_pessoal'])

        # Adicionar todos os campos calculados
        dados_peticao['valor_liquido'] = format_brl(renda - parcela_pessoal - total_consignado)
        dados_peticao['diario'] = format_brl((renda - parcela_pessoal - total_consignado) / 30)
        dados_peticao['comprometimento'] = format_brl(parcela_pessoal + total_consignado)
        dados_peticao['total_emprestimo'] = format_brl(total_emprestimo_geral)
        dados_peticao['def_emprestimos'] = format_brl(def_emprestimos)
        dados_peticao['parcela_pessoal_atual'] = format_brl(parcela_pessoal_atual)
        dados_peticao['dif_bacen'] = format_brl(dif_bacen)
        dados_peticao['vlr_total_emprestimo1'] = format_brl(vlr_total_emprestimo1)
        dados_peticao['vlr_total_emprestimo2'] = format_brl(vlr_total_emprestimo2)
        dados_peticao['org_bacen'] = format_brl(org_bacen)
        dados_peticao['org_div'] = format_brl(org_div)
        dados_peticao['total_dobro'] = format_brl(total_dobro)
        dados_peticao['total_dobro_geral'] = format_brl(total_dobro_geral)
        dados_peticao['valor_causa'] = format_brl(valor_causa)
        dados_peticao['dadovalorcausa'] = format_brl(dadovalorcausa)
        dados_peticao['comprometimento_renda'] = format_brl(comprometimento_renda)
        dados_peticao['renda_atual'] = format_brl(renda_atual)
        dados_peticao['comprometimento_porcentagem'] = format_brl(comprometimento_porcentagem)
        dados_peticao['total_emprestimo_bacen'] = format_brl(total_emprestimo_bacen)
        
        # DEBUG: Log final dos dados
        app.logger.debug(f"Dados completos para geração: {dados_peticao}")
        
        # Usar a nova função gerar_documento que determina o template correto
        documento = gerar_documento(dados_peticao, num_emprestimos, foro)
        app.logger.debug(f"Documento gerado com sucesso para foro: {foro}")
        
        return send_file(
            documento,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            download_name=f"peticao_{foro}_{num_emprestimos}_{datetime.now().strftime('%Y%m%d')}.docx",
            as_attachment=True
        )
    
    except ValueError as e:
        app.logger.error(f"Erro de validação: {str(e)}")
        return f"Erro nos dados fornecidos: {str(e)}", 400
    except FileNotFoundError as e:
        app.logger.error(f"Erro de arquivo: {str(e)}")
        return str(e), 404
    except Exception as e:
        app.logger.error(f"Erro inesperado: {str(e)}")
        import traceback
        app.logger.error(traceback.format_exc())
        return abort(500, "Erro interno ao processar a solicitação")

def flatten_dict(d, parent_key='', sep='_'):
    stack = [(d, parent_key)]
    items = []
    while stack:
        current_dict, current_key = stack.pop()
        for k, v in current_dict.items():
            new_key = f"{current_key}{sep}{k}" if current_key else k
            if isinstance(v, dict):
                stack.append((v, new_key))
            elif isinstance(v, list):
                for i, item in enumerate(v):
                    stack.append((item, f"{new_key}{sep}{i}"))
            else:
                items.append((new_key, v))
    return dict(items)

@app.route('/documentos/gerar', methods=['POST'])
@login_required
def gerar_documentos():
    app.logger.debug("Acessando rota /documentos/gerar")
    selecionados = request.form.getlist('documentos')
    if not selecionados:
        app.logger.error("Nenhum documento selecionado")
        abort(400, 'Nenhum documento selecionado.')

    # Processar dados do formulário
    placeholders = {}
    for key in request.form.keys():
        if key != 'documentos':
            value = request.form.get(key, '')
            
            # Processar campos especiais (listas)
            if key in ['conjunto_probatorio', 'fontes_renda']:
                placeholders[key] = formatar_lista_selecionados(value, "json")
            elif key in ['foro', 'tipo_peticao', 'possui_emprestimos']:
                placeholders[key] = MAPEAMENTO_VALORES.get(value, value)
            # ADICIONAR O CAMPO BENEFICIOS
            elif key == 'beneficios':
                placeholders[key] = value  # Já vem formatado do frontend
                placeholders['BENEFICIOS'] = value  # Também adicionar versão em maiúsculas
            else:
                placeholders[key] = value

    # Determinar número de empréstimos se for uma petição
    num_emprestimos = 1  # padrão
    if 'modelo_peticao' in request.form:
        try:
            num_emprestimos = int(request.form.get('modelo_peticao'))
            if num_emprestimos not in [1, 2, 3]:
                num_emprestimos = 1
        except (ValueError, TypeError):
            num_emprestimos = 1

    # Determinar foro
    foro = request.form.get('foro', 'autor').lower()
    app.logger.debug(f"Foro selecionado para documentos: {foro}")

    # Filtrar documentos selecionados válidos
    documentos_validos = []
    for doc in selecionados:
        if doc in VALID_TEMPLATES:
            documentos_validos.append(doc)
        else:
            # Verificar se é uma petição que precisa ser mapeada
            if doc.startswith('peticao_'):
                # Determinar o template correto baseado no foro e número de empréstimos
                try:
                    template_name = determinar_template_peticao(foro, num_emprestimos)
                    if template_name in VALID_TEMPLATES:
                        documentos_validos.append(template_name)
                        app.logger.debug(f"Petição mapeada para template: {template_name}")
                    else:
                        app.logger.warning(f"Template {template_name} não é válido")
                except FileNotFoundError as e:
                    app.logger.warning(str(e))
            else:
                app.logger.warning(f"Documento {doc} não é válido")

    if not documentos_validos:
        app.logger.error("Nenhum documento válido selecionado")
        abort(400, 'Nenhum documento válido selecionado.')

    # Dicionário de trechos que devem estar em negrito
    bold_sections = {
        "CONTRATADO: GUILHERME ESTEVES DOS SANTOS MORAES": True,
    }

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zipf:
        for doc_tipo in documentos_validos:
            docx_path = os.path.join('modelos', f"{doc_tipo}.docx")
            if not os.path.exists(docx_path):
                app.logger.warning(f"Template {doc_tipo}.docx não encontrado.")
                continue

            # Preencher DOCX
            doc = Document(docx_path)
            
            # Processar parágrafos
            for p in doc.paragraphs:
                # Substituir placeholders
                for chave, valor in placeholders.items():
                    if f'{{{{{chave}}}}}' in p.text:
                        p.text = p.text.replace(f'{{{{{chave}}}}}', bleach.clean(str(valor)))
                
                # Aplicar negrito nos trechos específicos
                for text, should_bold in bold_sections.items():
                    if text in p.text and should_bold:
                        for run in p.runs:
                            if text in run.text:
                                run.bold = True
            
            # Processar tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Substituir placeholders
                        for chave, valor in placeholders.items():
                            if f'{{{{{chave}}}}}' in cell.text:
                                cell.text = cell.text.replace(f'{{{{{chave}}}}}', bleach.clean(str(valor)))
                        
                        # Aplicar negrito nos trechos específicos
                        for text, should_bold in bold_sections.items():
                            if text in cell.text and should_bold:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        if text in run.text:
                                            run.bold = True

            # Salvar DOCX em memória e adicionar no ZIP
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            zipf.writestr(f"{doc_tipo}.docx", output.read())

    zip_buffer.seek(0)
    app.logger.debug("Arquivo ZIP gerado com sucesso")
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        download_name='documentos.zip',
        as_attachment=True
    )

@app.route('/gerar-peticao-completa', methods=['POST'])
@login_required
def gerar_peticao_completa():
    """Rota para gerar petições específicas baseadas no foro e número de empréstimos"""
    try:
        app.logger.debug("Iniciando geração de petição completa")
        
        # Obter dados do formulário
        foro = request.form.get('foro', 'autor').lower()
        
        # Determinar número de empréstimos
        num_emprestimos = 1
        if 'modelo_peticao' in request.form:
            try:
                num_emprestimos = int(request.form.get('modelo_peticao'))
                if num_emprestimos not in [1, 2, 3]:
                    num_emprestimos = 1
            except (ValueError, TypeError):
                num_emprestimos = 1
        
        # Determinar qual template usar
        template_name = determinar_template_peticao(foro, num_emprestimos)
        
        # Processar dados do formulário
        placeholders = {}
        for key in request.form.keys():
            value = request.form.get(key, '')
            
            # Processar campos especiais (listas)
            if key in ['conjunto_probatorio', 'fontes_renda']:
                placeholders[key] = formatar_lista_selecionados(value, "json")
            elif key in ['foro', 'tipo_peticao', 'possui_emprestimos']:
                placeholders[key] = MAPEAMENTO_VALORES.get(value, value)
            # ADICIONAR O CAMPO BENEFICIOS
            elif key == 'beneficios':
                placeholders[key] = value
                placeholders['BENEFICIOS'] = value
            else:
                placeholders[key] = value
        
        # Gerar documento
        docx_path = os.path.join('modelos', f"{template_name}.docx")
        if not os.path.exists(docx_path):
            return f"Template {template_name}.docx não encontrado", 404
        
        doc = Document(docx_path)
        
        # Dicionário de trechos que devem estar em negrito
        bold_sections = {
            "CONTRATADO: GUILHERME ESTEVES DOS SANTOS MORAES": True,
        }
        
        # Processar parágrafos
        for p in doc.paragraphs:
            # Substituir placeholders
            for chave, valor in placeholders.items():
                if f'{{{{{chave}}}}}' in p.text:
                    p.text = p.text.replace(f'{{{{{chave}}}}}', bleach.clean(str(valor)))
            
            # Aplicar negrito nos trechos específicos
            for text, should_bold in bold_sections.items():
                if text in p.text and should_bold:
                    for run in p.runs:
                        if text in run.text:
                            run.bold = True
        
        # Processar tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Substituir placeholders
                    for chave, valor in placeholders.items():
                        if f'{{{{{chave}}}}}' in cell.text:
                            cell.text = cell.text.replace(f'{{{{{chave}}}}}', bleach.clean(str(valor)))
                    
                    # Aplicar negrito nos trechos específicos
                    for text, should_bold in bold_sections.items():
                        if text in cell.text and should_bold:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if text in run.text:
                                        run.bold = True
        
        # Salvar e retornar documento
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            download_name=f"{template_name}_{datetime.now().strftime('%Y%m%d')}.docx",
            as_attachment=True
        )
    
    except Exception as e:
        app.logger.error(f"Erro inesperado ao gerar petição completa: {str(e)}")
        return abort(500, "Erro interno ao processar a solicitação")

@app.route('/numero_contrato')
@login_required
def numero_contrato():
    return render_template('numero_contrato.html')

if __name__ == '__main__':
    app.run(debug=True, port=5000)